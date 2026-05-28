import os
import re
import time
import tempfile

from dotenv import load_dotenv
from typing import TypedDict, List, Any, Dict

from langchain_groq import ChatGroq
from langgraph.graph import StateGraph, END, START
from langchain_core.messages import HumanMessage
from langchain_community.document_loaders import Docx2txtLoader
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.prompts import PromptTemplate
from langchain_core.output_parsers import PydanticOutputParser, StrOutputParser
from pydantic import BaseModel, Field
from deep_translator import GoogleTranslator
from gtts import gTTS

from dictio import HINGLISH_TO_DEVANAGARI

load_dotenv()


def hinglish_to_devanagari(text):
    """Convert Hinglish (romanized Hindi) text to Devanagari script
    using the local dictionary, so Hindi TTS pronounces words correctly.
    English words not in the dictionary are kept as-is."""
    # Build a set of phrase lengths present in the dictionary for multi-word matching
    max_phrase_len = max(
        (len(k.split()) for k in HINGLISH_TO_DEVANAGARI), default=1
    )
    words = text.split()
    result = []
    i = 0
    while i < len(words):
        matched = False
        # Try longest phrase first, then shorter
        for n in range(min(max_phrase_len, len(words) - i), 0, -1):
            phrase = " ".join(words[i:i + n])
            lookup = phrase.lower().strip(".,!?;:\"'()")
            if lookup in HINGLISH_TO_DEVANAGARI:
                result.append(HINGLISH_TO_DEVANAGARI[lookup])
                i += n
                matched = True
                break
        if not matched:
            # Try single word after stripping punctuation
            raw = words[i]
            stripped = raw.lower().strip(".,!?;:\"'()")
            if stripped in HINGLISH_TO_DEVANAGARI:
                # Preserve any trailing punctuation
                converted = HINGLISH_TO_DEVANAGARI[stripped]
                trailing = ""
                for ch in reversed(raw):
                    if ch in ".,!?;:\"'()":
                        trailing = ch + trailing
                    else:
                        break
                result.append(converted + trailing)
            else:
                result.append(raw)  # keep English/unknown words as-is
            i += 1
    return " ".join(result)

LANGUAGE_CODES = {
    "hindi": "hi",
    "hinglish": "hi",
    "gujarati": "gu",
    "marathi": "mr",
    "bengali": "bn",
    "tamil": "ta",
    "telugu": "te",
    "urdu": "ur",
    "english": "en",
}

model_text = ChatGroq(
    model="meta-llama/llama-4-scout-17b-16e-instruct",
    # model="llama-3.3-70b-versatile",
    api_key=os.getenv("GROQ_API_KEY"),
    temperature=0.5,
)
# model_text = ChatHuggingFace(llm=llm)

def _invoke_with_retry(chain, inputs, max_retries=3, delay=1):
    """Invoke a LangChain chain, retrying on any exception (including JSON parse failures)."""
    import traceback as _tb
    last_exc = None
    for attempt in range(max_retries):
        try:
            return chain.invoke(inputs)
        except Exception as e:
            last_exc = e
            print(f"[_invoke_with_retry] attempt {attempt+1}/{max_retries} failed: {type(e).__name__}: {e}")
            _tb.print_exc()
            if attempt < max_retries - 1:
                time.sleep(delay)
    raise last_exc

def _extract_image_content(image_path: str) -> str:
    """Use the vision-capable Groq model to extract all text and visual content from an image."""
    import base64
    ext = image_path.split(".")[-1].lower()
    mime_type = "image/png" if ext == "png" else "image/jpeg"
    with open(image_path, "rb") as f:
        image_b64 = base64.b64encode(f.read()).decode()
    message = HumanMessage(content=[
        {
            "type": "text",
            "text": (
                "This is an educational slide or document image. "
                "Extract ALL text content exactly as it appears — headings, bullet points, "
                "body text, labels, captions, equations, and numbers. "
                "Also describe any diagrams, charts, tables, or visual elements in detail. "
                "Present everything in a structured, readable format."
            ),
        },
        {
            "type": "image_url",
            "image_url": {"url": f"data:{mime_type};base64,{image_b64}"},
        },
    ])
    result = model_text.invoke([message])
    return result.content


def ingestion(path):
    file_type = path.split(".")[-1].lower()
    if file_type == "pdf":
        import fitz
        doc_pdf = fitz.open(path)
        documents = []
        for page_num in range(len(doc_pdf)):
            page = doc_pdf[page_num]
            # sort=True reads blocks in visual top-to-bottom, left-to-right order,
            # which prevents the reversed character-spaced text artifact
            text = page.get_text("text", sort=True)
            documents.append(Document(
                page_content=text,
                metadata={"page": page_num, "source": path}
            ))
        doc_pdf.close()
    elif file_type == "docx":
        loader = Docx2txtLoader(path)
        documents = loader.load()
    elif file_type in ("png", "jpg", "jpeg"):
        extracted_text = _extract_image_content(path)
        documents = [Document(
            page_content=extracted_text,
            metadata={"page": 0, "source": path}
        )]
    else:
        raise ValueError("Unsupported file format")
    return documents

#=======================     CLEANING ====================================================

def cleaning(documents):
    cleaned_docs = []
    for doc in documents:
        text = doc.page_content
        text = re.sub(r'Page\s+\d+', ' ', text)
        text = re.sub(r'^\s*\d+\s*$', ' ', text, flags=re.MULTILINE)
        text = re.sub(r'[^\x00-\x7F]+', ' ', text)
        text = re.sub(r'\n+', '\n', text)
        text = re.sub(r'(?<!\n)\n(?!\n)', ' ', text)
        text = re.sub(r'\s+', ' ', text)
        # Fix char-spaced reversed text: a PDF artifact where individual characters
        # are space-separated and stored in reverse visual order (e.g. "m e t I d o o F")
        text = re.sub(
            r'(?<![A-Za-z])(?:[A-Za-z] ){4,}[A-Za-z](?![A-Za-z])',
            lambda m: m.group(0).replace(' ', '')[::-1],
            text
        )
        text = text.strip()
        cleaned_docs.append(
            Document(
                page_content=text,
                metadata=doc.metadata
            )
        )
    return cleaned_docs

#=======================     PREPROCESS =========================

def preprocess(path):
    docs = ingestion(path=path)
    clean_text = cleaning(docs)
    splitter = RecursiveCharacterTextSplitter(chunk_size=3500, chunk_overlap=100)
    splitted_doc = splitter.split_documents(clean_text)
    return splitted_doc


def data_extraction(preprocessed_data):
    new = []
    for i, data in enumerate(preprocessed_data):
        schema = {
            "page_no": data.metadata.get("page", i),  # fallback to chunk index for DOCX
            "doc": data.page_content
        }
        new.append(schema)
    return new

class state(TypedDict):
    pages:List[int]
    all_documents:List[Dict]
    extracted_docs:List[Dict]
    language:str
    explained_list:List[str]
    audio_list: List[str]
    task:str
    maths_explained:List[Any]
    
def extract_particular(state):
    page_no_list=state["pages"]
    all_documents=state["all_documents"]
    extracted_list=[]
    for page in all_documents:
        if page["page_no"] in page_no_list:
            ans=f"THIS CONTENT IS OF PAGE NUMBER {page['page_no']} \n {page['doc']} "
            extracted_list.append(ans)
    return {"extracted_docs":extracted_list}   




parser=StrOutputParser()


class TheoryConceptUsed(BaseModel):
    name: str = Field(description="Name of the concept or key term — e.g. 'Photosynthesis', 'Democracy', 'Evaporation'")
    concept_used_original: str = Field(description="Description of the concept exactly as it appears on the slide, in the slide's original language — do NOT translate. If the slide says it in Hindi, write in Hindi. If English, write in English. This is the raw original wording from the slide for this concept.")
    concept_used_explanation: str = Field(description="STRICTLY in the user's chosen target language ONLY. Simple friendly explanation of this concept. Casual warm tone, like explaining to a friend. 2-4 sentences. NEVER write this in the slide's original language — ALWAYS in the target language.")


class TheoryContentPoint(BaseModel):
    content_original: str = Field(description="The original point/fact from the slide in the slide's original language — do NOT translate. Copy the exact point as it appears on the slide, slightly expanded for clarity.")
    content_original_translation: str = Field(description="STRICTLY in the user's chosen target language ONLY. Translate this point and explain it briefly — 2-3 sentences max. Cover exactly what the point says, slightly unpacked. Do NOT add extra context, examples, or tangents beyond what the point itself contains. NEVER write this in the slide's original language — ALWAYS in the target language.")


class SlideExplanation(BaseModel):
    slide_summary: str = Field(description="1-2 line summary of what this slide is about, in the slide's original language")
    concepts_used: List[TheoryConceptUsed] = Field(description="List of key concepts/terms from the slide with their original description and translated explanation")
    content_points: List[TheoryContentPoint] = Field(description="The main theory content broken into individual points — each with original text and translated explanation")
    real_life_example: str = Field(description="A simple real life analogy or example in user's chosen language that connects the slide topic to something the student sees or experiences in daily life")


class TheoryResponseModel(BaseModel):
    slides: List[SlideExplanation]


theory_parser = PydanticOutputParser(pydantic_object=TheoryResponseModel)


def _classify_theory_content(content: str) -> str:
    """Classify whether the content is a story/narrative or an academic slide.
    Returns 'story' or 'academic'."""
    classify_prompt = PromptTemplate(
        template="""Read the following content carefully and classify it into exactly ONE of these two categories:

CLASSIFY AS "story" IF the content is:
- A fictional or real narrative / story / tale (folk tale, short story, novel chapter, fable, myth, legend)
- A biographical or autobiographical account (life events narrated in sequence)
- A historical narrative told in story form (events described with characters and plot)
- A passage from a language/literature textbook (English lesson, Hindi lesson, etc.)
- Any prose or poetry with characters, events, emotions, conflict, or resolution
- A news report or article written in narrative form

CLASSIFY AS "academic" IF the content is:
- A conceptual slide with definitions, facts, bullet points, or scientific topics
- Science, Geography, Civics, Economics, History (data/facts-focused, not narrative)
- A slide with headings, sub-headings, diagrams described, tables, or lists
- Technical or factual content about a subject (e.g., how photosynthesis works, types of soil)
- Content that reads like a textbook chapter explaining a concept, NOT narrating events

CONTENT:
{content}

Reply with EXACTLY one word: story OR academic
Nothing else. Just the single word.""",
        input_variables=["content"]
    )
    chain = classify_prompt | model_text
    result = chain.invoke({"content": content[:2000]})  # classify on first 2000 chars for speed
    answer = result.content.strip().lower()
    return "story" if "story" in answer else "academic"


def _story_explain(slide_text: str, language: str) -> TheoryResponseModel:
    """Generate rich story/narrative explanation mapped to the TheoryResponseModel schema."""
    prompt = PromptTemplate(
        template="""
You are a warm, engaging literature teacher explaining a story or narrative passage to a student.
Your job is to help the student truly understand WHAT happened, WHO the characters are, WHY events matter,
and WHAT lesson or theme the story carries — all in the student's chosen language.

Here is the story/passage content:
{slide_text}

Explain this in {language}.

---

## ⚠️ ABSOLUTE LANGUAGE RULE — READ THIS FIRST ⚠️
The following fields MUST be written ENTIRELY and ONLY in **{language}**:
  - concept_used_explanation  (character/element explanation in user's language)
  - content_original_translation  (event explanation in user's language)
  - real_life_example  (theme/moral in user's language)

NO EXCEPTIONS. Even if the passage is in English, these fields MUST be in {language}.
If {language} is "hinglish", write in casual romanized Hindi mixed with common English words.
If {language} is "hindi", write in Devanagari script Hindi.
If {language} is "english", write in simple clear English.
NEVER write the explanation fields in the slide/passage's original language.

---

## SCHEMA MAPPING — READ CAREFULLY

You MUST return a JSON in the TheoryResponseModel schema. Here is how to map story elements to that schema:

### slide_summary:
- 2-3 sentence overview of what this passage/chapter is about.
- Write in the SAME LANGUAGE as the original passage (do NOT translate).
- Cover: Who is the story about? What happens? What is the overall mood or setting?

---

### concepts_used → USE THIS FOR: Characters + important narrative elements

For EACH major character or key narrative element (setting, symbol, conflict):

**name**: The character's name (e.g., "Della", "Jim", "The Gift of the Magi") OR element name (e.g., "Central Conflict", "Setting")

**concept_used_original**: Copy the key description of this character/element directly from the original passage — in the passage's original language. Do NOT translate. Quote or paraphrase what the text actually says about them. 2-3 sentences from the text.

**concept_used_explanation**: ⚠️ MUST BE 100% IN {language}.
Explain who this character is OR what this element means for the story.
- For a character: Who are they? What is their role (protagonist/antagonist/supporting)? What makes them important to the story? What do they want or fear? What are they like as a person?
- For a story element (conflict, setting, symbol): What is it? Why does it matter to the story?
- Length: 4-6 sentences. Warm, engaging tone. Not too short — the student should truly understand this character after reading.
- Examples: "Della ek bahut pyaari aur caring wife hai jo apne pati Jim se bahut zyada pyaar karti hai. Woh itni garib hai ki unke paas Christmas gift ke liye kuch nahi hai. Iska bohot important role hai kyunki poori story ussi ki problem ke baare mein hai — woh kya karega apne pati ke liye gift kharidne ke liye?"

Include at least 2-4 entries: all main characters + 1 entry for the setting/time if relevant.

---

### content_points → USE THIS FOR: Key events / plot points in sequence

Break the story into its MAIN EVENTS in chronological order. Each event = one content_point.
Aim for 5-8 events that together tell the complete story arc.

**content_original**: The key event as it appears/is described in the original passage language. 
Copy the relevant lines or paraphrase the event in the original language. 2-3 sentences from the text.
Do NOT translate.

**content_original_translation**: ⚠️ MUST BE 100% IN {language}.
Explain what happened in this event and WHY it matters to the story.
- WHAT: Describe exactly what happened clearly and vividly — as if retelling to someone who hasn't read it.
- WHY IT MATTERS: Explain how this event moves the story forward or reveals something important about a character or theme.
- EMOTION: If there's an emotional moment, describe what the character feels and why the reader should care.
- LENGTH: 3-5 sentences minimum per event. NOT TOO SHORT. Each event explanation must be substantive and complete.
- DO NOT write one-liners. Each event deserves a proper explanation.
- Example: "Della ne apne lambe khoobsurat baalon ko ek shop mein 20 dollars mein bech diya. Yeh ek bohot emotional moment tha kyunki uske baal hi uski sabse badi khubsurti thi — Jim ko unse bahut pyaar tha. Lekin usne yeh kurbani is liye di kyunki woh Jim ke liye ek khaas gift — ek platinum watch chain — kharidna chahti thi. Iss ek action se pata chalta hai ki Della ka pyaar uske baalon se bhi zyada keemat rakhta tha."

---

### real_life_example → USE THIS FOR: Central theme + moral + real-life connection

Write ONE rich paragraph that covers:
1. The central THEME of the story (what big idea does it explore? — love, sacrifice, greed, courage, etc.)
2. The MORAL or LESSON (what does the student learn from this story?)
3. A REAL-LIFE CONNECTION (how does this theme appear in the student's own daily life?)
- ⚠️ MUST be in {language}.
- Length: 4-6 sentences. Rich, meaningful, thoughtful.
- Example: "Is kahaani ka central theme hai selfless love — matlab woh pyaar jo kuch bhi expect nahi karta. Della aur Jim dono ne apni-apni sabse keemat cheez kurbaan kar di sirf ek doosre ko khush karne ke liye. Iska lesson yeh hai ki sachcha pyaar sacrifice mein hota hai, gift ki value mein nahi. Asal life mein bhi aisa hota hai — jab koi apni neend chhodke tumhari takleef mein tumhare saath baitha rahe, yahi sachcha pyaar hai."

---

## ⚠️ LENGTH RULES — VERY IMPORTANT
- concept_used_explanation: MINIMUM 4 sentences. Must be rich and complete.
- content_original_translation: MINIMUM 3 sentences. Capture what + why + emotion.
- real_life_example: MINIMUM 4 sentences. Theme + moral + real-life connection.
- DO NOT write short, incomplete explanations. Quality over brevity.

---

## OUTPUT FORMAT

Return ONLY a strict JSON object. No extra text, no markdown, no explanation outside JSON.

{{
    "slides": [
        {{
            "slide_summary": "2-3 sentence overview in the passage's original language",
            "concepts_used": [
                {{
                    "name": "Character or Element Name",
                    "concept_used_original": "Key description from the original passage in the passage's original language",
                    "concept_used_explanation": "MUST BE IN {language} — 4-6 sentences about who this character is and why they matter"
                }}
            ],
            "content_points": [
                {{
                    "content_original": "Description of this event from the original passage in the passage's original language",
                    "content_original_translation": "MUST BE IN {language} — 3-5 sentences: what happened, why it matters, and the emotional weight"
                }}
            ],
            "real_life_example": "MUST BE IN {language} — 4-6 sentences: theme + moral + real-life connection"
        }}
    ]
}}

---

## FINAL CHECK BEFORE RETURNING:
1. Is EVERY concept_used_explanation in {language}? Each one minimum 4 sentences?
2. Is EVERY content_original_translation in {language}? Each one minimum 3 sentences?
3. Is real_life_example in {language}? Minimum 4 sentences covering theme + moral + real life?
4. Are there at least 5 content_points covering the full story arc?
5. Are there at least 2 characters explained in concepts_used?

Target Language: {language}

Passage Content:
{slide_text}

You MUST strictly follow the given schema:
{format_instructions}
""",
        input_variables=["slide_text", "language"],
        partial_variables={"format_instructions": theory_parser.get_format_instructions()}
    )
    chain = prompt | model_text | theory_parser
    return _invoke_with_retry(chain, {"slide_text": slide_text, "language": language})


def _academic_slide_explain(slide_text: str, language: str) -> TheoryResponseModel:
    """Generate explanation for academic/conceptual slide content."""
    prompt = PromptTemplate(
        template="""
You are explaining a theory slide to your friend — like a friend who knows the subject well,
talking to another friend who is studying Science, Social Studies, History, Geography,
Biology, Chemistry, Physics, or any other theoretical subject.
Your tone is casual, warm, and simple — like a WhatsApp voice note between friends.

Here is the slide content:
{slide_text}

Explain this in {language}.

---

## ⚠️ ABSOLUTE LANGUAGE RULE — READ THIS FIRST ⚠️
The following fields MUST be written ENTIRELY and ONLY in **{language}**:
  - concept_used_explanation
  - content_original_translation
  - real_life_example

NO EXCEPTIONS. Even if the slide is in English, these fields MUST still be in {language}.
If {language} is "hinglish", write in casual romanized Hindi mixed with common English words.
If {language} is "hindi", write in Devanagari script Hindi with common English technical terms kept as-is.
If {language} is "english", write in simple casual English.
EVERY SINGLE WORD of these three fields must be in {language}. Do NOT fall back to the slide's language.
If you write even one sentence of concept_used_explanation or content_original_translation in the wrong language, the output is REJECTED.

---

## STRICT RULES

### IMPORTANT — THIS IS A THEORY SUBJECT:
- This is NOT a mathematics slide. There are NO steps, NO calculations, NO arithmetic to walk through.
- Your job is to make the CONCEPT clear — what it means, why it happens, why it matters.

---

### slide_summary RULES:
- Write 1-2 lines summarizing what this slide is about.
- Must be in THE SAME LANGUAGE AS THE SLIDE CONTENT — do NOT translate into {language}.
- Clean readable sentence — not raw bullet points copied from the slide.
  ❌ Wrong: "Photosynthesis — definition, process, chlorophyll"
  ✅ Right: "This slide explains what photosynthesis is and how plants use sunlight to make their own food."

---

### concepts_used RULES:
- Identify ALL key concepts, terms, or vocabulary used in the slide.
- For EACH concept, provide:
  1. **name** — short name of the concept (e.g. "Photosynthesis", "Democracy")
  2. **concept_used_original** — copy the description/definition of this concept EXACTLY as it appears on the slide, in the slide's original language. Do NOT translate. If the slide doesn't define it, write a 1-2 line description in the slide's original language.
  3. **concept_used_explanation** — ⚠️ THIS MUST BE 100% IN {language}. Simple friendly explanation of this concept. Casual warm tone, like explaining to a friend. 2-4 sentences. Keep English technical terms (like "photosynthesis", "democracy") as-is, but ALL other words MUST be in {language}.
- Include at least 2-4 concepts per slide (more if the slide is content-heavy).
- Examples of tone for concept_used_explanation when language is hinglish:
    ✅ "Photosynthesis matlab woh process jisme plants sunlight use karke apna food banaate hain. Isme CO2 aur water se glucose banta hai."
    ❌ "Photosynthesis is the process by which plants make their food using sunlight." ← WRONG, this is English not hinglish
    ❌ "प्रकाश संश्लेषण वह प्रक्रिया है जिसमें पौधे सूर्य के प्रकाश का उपयोग करके भोजन बनाते हैं।" ← WRONG if language is hinglish

---

### content_points RULES (MOST IMPORTANT):
- Break the slide content into individual points/facts.
- Each point becomes one entry in the content_points list.
- For EACH point, provide:
  1. **content_original** — the original point from the slide, in the slide's original language. Copy the exact text, slightly expanded for clarity if needed. Do NOT translate.
  2. **content_original_translation** — ⚠️ THIS MUST BE 100% IN {language}. Translate this point and explain it briefly.
     - SCOPE: Explain ONLY what this specific point says — do NOT add examples, do NOT explain related concepts, do NOT pad with background information.
     - LENGTH: 2-3 sentences MAX per point. If the original point is a short fact, 1-2 sentences is enough.
     - WHAT to include: translate the point, then add at most one sentence saying what it means or why it matters — that's it.
     - Keep English technical/subject terms as-is — but ALL other words MUST be in {language}.
     - Short sentences. Simple words. No bookish language.
     - NEVER copy the original text as-is. Always rewrite in {language}.
     - ❌ DO NOT: write a paragraph about the broader topic
     - ❌ DO NOT: add real-life examples here (that goes in real_life_example only)
     - ❌ DO NOT: explain concepts that are already in concepts_used
- COVER EVERY SINGLE POINT on the slide — do not skip anything.
- If the slide has 8 points, create 8 content_points entries.
- Do NOT merge multiple points into one.
- Order content_points in the same order as they appear on the slide.
- Examples when language is hinglish:
    content_original: "Photosynthesis is the process by which green plants prepare their own food using sunlight."
    content_original_translation: "Photosynthesis ek aisi process hai jisme green plants sunlight use karke apna food khud banaate hain." ← SHORT, direct, covers just the point
    ❌ WRONG (too long): "Toh basically photosynthesis ek process hai... [3 more sentences explaining chlorophyll, CO2, glucose...]" ← That's too much for one point
    ❌ WRONG (just copying): "Photosynthesis is the process by which green plants prepare their own food using sunlight." ← Not translated

---

### real_life_example RULES:
- Give ONE strong real-life analogy or example that connects the slide topic to
  something the student sees, does, or experiences in daily life.
- ⚠️ MUST be in {language} — not in the slide's language.
- Should make the student say "ohh yeh toh main roz dekhta hoon!"
- Keep it to 2-3 sentences max.

---

## OUTPUT FORMAT

Return ONLY a strict JSON object. No extra text, no markdown, no explanation outside JSON.

{{
    "slides": [
        {{
            "slide_summary": "1-2 line summary in the slide's original language",
            "concepts_used": [
                {{
                    "name": "Concept Name",
                    "concept_used_original": "Exact description from the slide in the slide's original language",
                    "concept_used_explanation": "MUST BE IN {language} — Simple friendly explanation of this concept"
                }}
            ],
            "content_points": [
                {{
                    "content_original": "First point from the slide in the slide's original language",
                    "content_original_translation": "MUST BE IN {language} — 2-3 sentences MAX: translate the point, then at most one sentence on what it means. No extra padding."
                }},
                {{
                    "content_original": "Second point from the slide in the slide's original language",
                    "content_original_translation": "MUST BE IN {language} — 2-3 sentences MAX: translate the point, then at most one sentence on what it means. No extra padding."
                }}
            ],
            "real_life_example": "MUST BE IN {language} — 2-3 sentence real life analogy"
        }}
    ]
}}

---

## FINAL LANGUAGE CHECK BEFORE RETURNING:
Before you return the JSON, verify:
1. Is EVERY concept_used_explanation written in {language}? If not, rewrite it.
2. Is EVERY content_original_translation written in {language}? If not, rewrite it.
3. Is real_life_example written in {language}? If not, rewrite it.
4. ONLY slide_summary and concept_used_original and content_original should be in the slide's original language.
5. Everything else MUST be in {language}.

Target Language: {language}

Slide Content:
{slide_text}

Return the JSON object only. No extra text outside JSON.
You MUST strictly follow the given schema:
{format_instructions}
""",
        input_variables=["slide_text", "language"],
        partial_variables={"format_instructions": theory_parser.get_format_instructions()}
    )
    chain = prompt | model_text | theory_parser
    return _invoke_with_retry(chain, {"slide_text": slide_text, "language": language})


def theory_trans(slide_text_list, language):
    """Classify each page as story/narrative or academic slide, then explain accordingly."""
    explanation_list = []
    for slide_text in slide_text_list:
        content_type = _classify_theory_content(slide_text)
        if content_type == "story":
            result = _story_explain(slide_text, language)
        else:
            result = _academic_slide_explain(slide_text, language)
        explanation_list.append(result)
    return explanation_list
    
class Step(BaseModel):
    step_number: int
    title: str = Field(description="Short title of what this step does, e.g. 'Prime factorize 12' or 'Find common factors'")
    math_working: List[str] = Field(description="Line-by-line mathematical computation. Each string is one line of math, e.g. '12 = 2 × 2 × 3', 'HCF = 2 × 3 = 6'. Must contain ACTUAL numbers and operations, NOT descriptions.")
    simple_explanation: str = Field(description="Friendly explanation in user's chosen language of what happened in this step and why — narrated verbally as if speaking to the student, no symbols")


class ConceptUsed(BaseModel):
    name: str = Field(description="Name of the concept, theorem, formula, or term — e.g. 'Pythagoras Theorem', 'HCF', 'Congruent'")
    concept_used_original: str = Field(description="The concept/term exactly as it appears in the slide's original language — copy the exact wording from the slide. Do NOT translate. If the slide is in English, copy in English. If in Hindi, copy in Hindi. This is the raw original text for this concept from the slide.")
    explanation: str = Field(description="Simple friendly explanation of what this concept means and why it is used here, in the user's chosen language. Write as if explaining to a slow learner — no jargon, warm tone, 3-5 sentences.")


class MathQuestion(BaseModel):
    question_number: int
    original_question_text: str = Field(description="The exact problem statement being solved")
    original_question_text_explanation: str = Field(description="A friendly explanation of what the question is asking, in the user's chosen language. Break down the question sentence by sentence so the student understands what is being asked before solving. 3-5 warm sentences.")
    given_info: str = Field(description="What values/information are given in the problem — list them clearly")
    to_find: str = Field(description="What we need to calculate or find — state it concisely")
    to_find_explanation: str = Field(description="A friendly explanation of what exactly we need to find and why — in the user's chosen language. Help the student visualise the goal before starting. 2-4 warm sentences.")
    steps: List[Step]
    final_answer: str = Field(description="The complete final answer clearly stated with units")


class MathResponseModel(BaseModel):
    concepts_used: List[ConceptUsed] = Field(description="List of ALL mathematical concepts, theorems, formulas, and special terms used anywhere across all questions on this slide. Cover every formula used, every theorem applied, every term like 'congruent', 'parallel', 'HCF', 'prime factor' etc. Explain each one simply in the user's chosen language.")
    questions: List[MathQuestion]

# ── Maths Concept-Only Schema (when page has no numerical problems) ──

class MathConceptDetail(BaseModel):
    name: str = Field(description="Name of the concept, theorem, formula, or property")
    concept_used_original: str = Field(description="The concept/theorem/formula exactly as it appears in the slide's original language — copy the exact wording from the slide. Do NOT translate.")
    definition: str = Field(description="Formal definition or statement of the concept/theorem STRICTLY in user's chosen target language (e.g. hinglish, hindi, english). If language is hinglish, write in romanized Hindi NOT in English. 2-3 sentences.")
    simple_explanation: str = Field(description="Detailed friendly explanation STRICTLY in user's chosen target language — if language is hinglish, write in romanized Hindi NOT plain English. Explain as if to a 12-year-old. Cover what it means, why it matters, and give a relatable analogy. 5-8 sentences. No symbols — write in words only. This will be used for AUDIO narration.")
    formula_or_expression: str = Field(description="The mathematical formula or expression if applicable, e.g. 'a² + b² = c²', 'Area = l × b'. Write 'N/A' if no formula applies.")
    when_to_use: str = Field(description="When and where this concept is applied — practical scenarios and problem types, STRICTLY in user's chosen target language (if hinglish, write romanized Hindi NOT English). 2-3 sentences.")
    common_mistakes: str = Field(description="Common mistakes or confusions students have with this concept, STRICTLY in user's chosen target language (if hinglish, write romanized Hindi NOT English). 2-3 sentences.")

class MathConceptResponseModel(BaseModel):
    page_summary: str = Field(description="Brief summary of what this page/slide covers, STRICTLY in user's chosen target language (if hinglish, write romanized Hindi NOT English). 2-3 sentences.")
    concepts: List[MathConceptDetail] = Field(description="List of ALL concepts, theorems, formulas, properties, and definitions covered on this page. Be exhaustive — include every single concept mentioned.")

maths_concept_parser = PydanticOutputParser(pydantic_object=MathConceptResponseModel)

maths_parser=PydanticOutputParser(pydantic_object=MathResponseModel)

def _classify_maths_content(context):
    """Classify whether maths slide has numerical problems or only concepts/definitions."""
    classify_prompt = PromptTemplate(
        template="""Look at this mathematics slide content and classify it.

CLASSIFY AS "numerical" IF the slide contains:
- Actual math problems to solve (e.g., "Find HCF of 12 and 18", "Solve: 2x + 3 = 7")
- Worked examples with specific numbers
- Exercises or practice problems with numerical values
- Word problems with quantities/measurements

CLASSIFY AS "conceptual" IF the slide ONLY contains:
- Definitions of concepts/theorems (e.g., "A prime number is...")
- Theorem statements without worked examples (e.g., "Pythagoras theorem states that...")
- Properties or rules listed (e.g., "Properties of parallel lines...")
- Formulas listed without being applied to specific numbers
- Introductory/theoretical content with no problems to solve

SLIDE CONTENT:
{context}

Reply with EXACTLY one word: numerical OR conceptual
Nothing else. Just the single word.""",
        input_variables=["context"]
    )
    chain = classify_prompt | model_text
    result = chain.invoke({"context": context})
    answer = result.content.strip().lower()
    return "conceptual" if "conceptual" in answer else "numerical"

def maths_explain(slide_text_list, language):
    
    maths_explained=[]
    for context in slide_text_list:
        
        content_type = _classify_maths_content(context)
        
        if content_type == "conceptual":
            result = _maths_concept_explain(context, language)
            maths_explained.append({"type": "conceptual", "data": result})
        else:
            result = _maths_numerical_explain(context, language)
            maths_explained.append({"type": "numerical", "data": result})
        
    return maths_explained

def _maths_concept_explain(context, language):
    """Explain maths concepts/theorems when no numerical problems exist."""
    prompt = PromptTemplate(
        template="""
You are an expert mathematics tutor who teaches slow learners and students who are afraid of maths.
This slide contains mathematical concepts, theorems, definitions, or formulas — but NO numerical problems to solve.

Your job is to explain EVERY concept, theorem, formula, and property on this slide in a warm, detailed, and
student-friendly way, as if you are sitting next to a 12-year-old and helping them truly understand each idea.

---

## INSTRUCTIONS

### page_summary:
- Summarize what this page/slide is about in {language}.
- 2-3 sentences that give an overview of the topics covered.

### For EACH concept on the page, provide:

#### name:
- The standard name of the concept/theorem/formula (e.g., "Pythagoras Theorem", "Prime Numbers", "Commutative Property")

#### concept_used_original:
- Copy the EXACT text from the slide for this concept — in the slide's original language.
- Do NOT translate. Just copy verbatim.

#### definition:
- The formal mathematical definition or statement in {language}.
- 2-3 clear sentences.

#### simple_explanation:
- This is the MOST IMPORTANT field — it will be used for AUDIO narration.
- Explain in {language} as if talking to a scared 12-year-old face-to-face.
- Cover: What is it? Why does it matter? Give a real-life analogy or relatable example.
- 5-8 sentences, warm and friendly tone.
- NEVER use mathematical symbols — write everything in words (e.g., "a square plus b square equals c square" not "a² + b² = c²").

#### formula_or_expression:
- The mathematical formula/expression if one exists (e.g., "a² + b² = c²", "Area = πr²").
- Write "N/A" if no formula applies.

#### when_to_use:
- When is this concept used? What types of problems require it?
- In {language}, 2-3 sentences with practical scenarios.

#### common_mistakes:
- What do students commonly get wrong or confused about?
- In {language}, 2-3 sentences.

---

## ⚠️ ABSOLUTE LANGUAGE RULE — READ THIS FIRST ⚠️
The following fields MUST be written ENTIRELY and ONLY in **{language}**:
  - page_summary
  - definition
  - simple_explanation
  - when_to_use
  - common_mistakes

NO EXCEPTIONS. Even if the slide is in English or any other language, these fields MUST still be in {language}.
If {language} is "hinglish", write in casual romanized Hindi mixed with common English words — meaning HINDI SENTENCES written in ROMAN SCRIPT with English technical terms kept as-is. This is NOT English. The sentence structure and grammar must be HINDI.
If {language} is "hindi", write in Devanagari script Hindi with common English technical terms kept as-is.
If {language} is "english", write in simple casual English.
EVERY SINGLE WORD of these fields must be in {language}. Do NOT fall back to the slide's language or to plain English.
If you write even one sentence of these fields in the wrong language, the output is REJECTED.

---

## HINGLISH EXAMPLES (use these as reference when {language} is "hinglish"):

### page_summary example:
  ✅ "Yeh page Pythagoras Theorem ke baare mein hai. Isme bataya gaya hai ki right-angled triangle ke sides ka kya relation hota hai."
  ❌ "This page is about Pythagoras Theorem. It explains the relationship between sides of a right-angled triangle." ← WRONG, this is plain English

### definition example:
  ✅ "Pythagoras Theorem kehta hai ki ek right-angled triangle mein, hypotenuse ka square baaki dono sides ke squares ke sum ke barabar hota hai."
  ❌ "Pythagoras Theorem states that in a right-angled triangle, the square of the hypotenuse equals the sum of the squares of the other two sides." ← WRONG, this is plain English

### simple_explanation example:
  ✅ "Dekho, Pythagoras Theorem bahut simple concept hai. Socho tumhare paas ek right angle wala triangle hai, matlab jisme ek angle 90 degrees ka hai. Ab isme sabse lambi side ko hypotenuse bolte hain. Toh yeh theorem kehta hai ki agar tum hypotenuse ko square karo, toh woh baaki dono sides ke squares ka total hoga. Jaise ek ladder deewar pe lagao — ladder ki length, deewar ki height, aur zameen pe distance, teeno ka ek fixed relation hota hai. Isliye yeh theorem real life mein bahut kaam aata hai."
  ❌ "Pythagoras Theorem is a simple concept. Consider a right-angled triangle where one angle is 90 degrees..." ← WRONG, this is plain English, NOT hinglish

### when_to_use example:
  ✅ "Jab bhi koi right-angled triangle mile aur tumhe koi missing side nikalni ho, tab Pythagoras Theorem use karo. Construction aur distance problems mein yeh bahut kaam aata hai."
  ❌ "Use Pythagoras Theorem when you need to find a missing side..." ← WRONG, plain English

### common_mistakes example:
  ✅ "Bahut saare students hypotenuse ki jagah koi aur side square kar dete hain. Yaad rakho hypotenuse HAMESHA sabse lambi side hoti hai aur woh 90 degree angle ke saamne hoti hai."
  ❌ "Many students square the wrong side instead of the hypotenuse..." ← WRONG, plain English

⚠️ KEY POINT: Hinglish means the GRAMMAR and SENTENCE STRUCTURE is Hindi, written in Roman/Latin script. Only technical terms (like theorem, triangle, hypotenuse, square) stay in English. Words like "is", "the", "of", "and", "it" should be replaced with their Hindi equivalents (hai, ka/ke/ki, aur, yeh) written in Roman script.

---

## IMPORTANT RULES
- Be EXHAUSTIVE — include EVERY concept, theorem, formula, property, and definition on the page.
- Do NOT skip anything, even if it seems minor.
- Do NOT invent numerical problems. Focus purely on explaining the concepts.
- All explanations must be in {language}.

---

## FINAL LANGUAGE CHECK BEFORE RETURNING:
Before you return the JSON, verify:
1. Is page_summary written in {language}? If not, rewrite it.
2. Is EVERY definition written in {language}? If not, rewrite it.
3. Is EVERY simple_explanation written in {language}? If not, rewrite it.
4. Is EVERY when_to_use written in {language}? If not, rewrite it.
5. Is EVERY common_mistakes written in {language}? If not, rewrite it.
6. ONLY concept_used_original and formula_or_expression should be in the slide's original language.
7. Everything else MUST be in {language}.
8. If {language} is "hinglish" — go back and CHECK: are the fields actually in romanized Hindi with English technical terms? Or did you accidentally write in plain English? If plain English, REWRITE in hinglish NOW.

---

Slide Content:
{context}

You MUST strictly follow the given schema:
{format_instructions}
""",
        input_variables=["language", "context"],
        partial_variables={"format_instructions": maths_concept_parser.get_format_instructions()}
    )
    chain = prompt | model_text | maths_concept_parser
    return _invoke_with_retry(chain, {"language": language, "context": context})

def _maths_numerical_explain(context, language):
    """Explain maths content that has numerical problems (original behavior)."""
    prompt = PromptTemplate(
    template="""
You are an expert mathematics tutor who teaches slow learners and students who are afraid of maths.
Your job is to explain the given mathematics slide content step by step with ACTUAL CALCULATIONS,
as if you are sitting next to a 12-year-old student and solving problems on paper together.

---

## CRITICAL RULE — ALWAYS SHOW REAL MATH

There are TWO types of maths slides:

### TYPE A: Slide has actual problems with numbers (e.g., "Find HCF of 12 and 18")
→ Solve the problem step by step with full arithmetic.

### TYPE B: Slide only has a concept/definition/method (e.g., "HCF is the highest common factor...")
→ You MUST INVENT a concrete numerical example and solve it fully.
   Example: If the slide defines HCF and its methods, you create:
   "Let us find HCF of 12 and 18 using Prime Factorization method"
   and solve it completely with real numbers.

⚠️ NEVER return steps that are just descriptions like "Express as product of prime factors".
   EVERY step MUST have actual numbers and calculations in math_working.

---

## SCHEMA RULES

### concepts_used (FILL THIS FIRST — covers the WHOLE page):
- List EVERY concept, theorem, formula, or special term used anywhere across all questions on this page.
- Include: formulas (Area = l × b), theorems (Pythagoras, Euclid), terms (HCF, LCM, prime factor,
  congruent, parallel, concurrent, rational, irrational, quadratic, etc.)
- For each concept:
  - name: the exact name/term (e.g. "Pythagoras Theorem", "Prime Factorization", "Congruent Triangles")
  - concept_used_original: copy the exact wording of this concept/term as it appears in the slide, in the slide's original language — do NOT translate. This is purely a verbatim extract from the slide text.
  - explanation: explain in {language} as if to a slow learner who has never heard this before.
    - What is it? What does it mean?
    - Why do we use it here?
    - Give a simple real-world example or analogy if possible.
    - 3-5 sentences, warm and friendly tone, no jargon.
    - NEVER use mathematical symbols in explanation — write everything in words.

---

### original_question_text:
- If the slide has an actual problem → copy it exactly.
- If the slide is only a concept → write: "Example: [your invented problem statement]"

### original_question_text_explanation:
- In {language}, explain what the question is asking in simple words.
- Break it down sentence by sentence — what information are we given, what is confusing, what do we need to do?
- 3-5 warm sentences. Tone: like a friend reading the question aloud and making sure you understand it.
- No symbols — write everything in words.
- Example: "Yeh question humse pooch raha hai ki do numbers ka HCF kya hoga. Pehle unhone humhe do numbers diye hain — 12 aur 18. Ab humhe dhundhna hai ki dono numbers mein se sabse bada number kaun sa hai jo dono ko exactly divide kar sake."

### given_info:
- List the actual numbers/values. e.g., "Numbers: 12, 18"

### to_find:
- What we are calculating. e.g., "HCF of 12 and 18" — concise label only.

### to_find_explanation:
- In {language}, explain exactly what we are trying to find and why it matters.
- Help the student picture the goal before we start solving.
- 2-4 warm sentences, no symbols.
- Example: "Humhe HCF nikalna hai, matlab Highest Common Factor. Iska matlab hai woh sabse bada number jo dono numbers — 12 aur 18 — ko bina kuch remainder ke divide kar sake. Agar hum yeh jaan lein, toh hum easily fractions simplify kar sakte hain ya koi bhi sharing wali problem solve kar sakte hain."

---

### title (per step):
- Short action phrase. e.g., "Prime factorize 12", "Identify common factors", "Calculate HCF"

### math_working (per step) — MOST IMPORTANT:
- Each item is ONE line of actual mathematical work.
- Must contain real numbers, real operations, real results.
- Examples of CORRECT math_working lines:
    "12 = 2 × 2 × 3"
    "18 = 2 × 3 × 3"
    "Common prime factors = 2, 3"
    "HCF = 2 × 3 = 6"
    "Area = length × breadth = 12 × 5 = 60 cm²"
    "x + 5 = 12"
    "x = 12 - 5 = 7"
- Examples of WRONG math_working (NEVER do this):
    ❌ "Express the number as product of prime factors"
    ❌ "Divide by smallest prime"
    ❌ "Apply the formula"
- Minimum 2-3 lines of math per step. Show every intermediate calculation.
- Use × for multiplication, ÷ for division, = for equals.

### simple_explanation (per step) — USED FOR AUDIO:
- In {language}, explain warmly like a patient friend teaching a slow learner face to face.
- Minimum 6-10 sentences per step.
- STRUCTURE exactly like this:
    1. Say WHAT this step is doing and WHY in context of the full problem.
    2. Verbally read out EVERY SINGLE line of math_working one line at a time in words:
       - "12 ÷ 2 = 6" → "pehle hum 12 ko 2 se divide karte hain, toh 12 divide by 2 barabar 6 aata hai"
       Do NOT skip ANY line. Do NOT group multiple lines.
    3. Connect it to the goal after all lines.
    4. End with an encouraging line like "dekha? kuch mushkil nahi tha!"
- NEVER use symbols ×, ÷, = — write in words: multiply, divide karke, barabar, plus, minus.

### final_answer:
- Complete answer with units. e.g., "HCF of 12 and 18 = 6"

---

## OUTPUT FORMAT

Return ONLY a strict JSON object. No markdown, no text outside JSON.

{{
    "concepts_used": [
        {{
            "name": "Prime Factorization",
            "concept_used_original": "it will includes all the concepts that is covered in this particulr page, what this concept is and why we use it here, in the slide's original language",
            "explanation": "Friendly explanation in {language} of what this concept is and why we use it here — no symbols, 3-5 sentences."
        }}
    ],
    "questions": [
        {{
            "question_number": 1,
            "original_question_text": "Find the HCF of 12 and 18",
            "original_question_text_explanation": "Friendly breakdown of what the question is asking in {language}, 3-5 sentences, no symbols.",
            "given_info": "Numbers: 12, 18",
            "to_find": "HCF of 12 and 18",
            "to_find_explanation": "Friendly explanation of what we need to find and why in {language}, 2-4 sentences, no symbols.",
            "steps": [
                {{
                    "step_number": 1,
                    "title": "Prime factorize 12",
                    "math_working": [
                        "12 ÷ 2 = 6",
                        "6 ÷ 2 = 3",
                        "3 ÷ 3 = 1",
                        "So, 12 = 2 × 2 × 3"
                    ],
                    "simple_explanation": "Verbally walk through EVERY math_working line in {language}. No symbols."
                }}
            ],
            "final_answer": "HCF of 12 and 18 = 6"
        }}
    ]
}}

---

Target Language: {language}

Slide Content:
{context}

Return the JSON object only. No extra text outside JSON.
You MUST strictly follow the given schema:
{format_instructions}
    """,
    input_variables=["language", "context"],
    partial_variables={"format_instructions": maths_parser.get_format_instructions()}
)
    chain = prompt | model_text | maths_parser
    return _invoke_with_retry(chain, {"language": language, "context": context})

def single_multi_page(state):
    language=state["language"]
    task=state["task"]
    slide_text_list=state["extracted_docs"]
    
    if task=="theory":
        ans=theory_trans(slide_text_list,language)
        
        return {"explained_list":ans}
        
    
    else:
        ans=maths_explain(slide_text_list,language)
        
        return {"maths_explained":ans}


import edge_tts
import asyncio
import tempfile

# ── Language-aware transition phrases for audio narration ──
TRANSITION_PHRASES = {
    "english": {
        "concepts_intro": "First, let us understand some important concepts from this page.",
        "concept_detail_intro": "Now let us understand each concept in detail.",
        "common_mistake_prefix": "One important thing to remember.",
        "real_life_prefix": "Let me explain with a real life example.",
        "numerical_concepts_intro": "First let us understand some important concepts used on this page.",
        "to_find_prefix": "We need to find:",
        "final_answer_prefix": "Final answer:",
    },
    "hindi": {
        "concepts_intro": "पहले कुछ ज़रूरी concepts समझ लेते हैं जो इस page पर हैं।",
        "concept_detail_intro": "अब हम एक एक concept को detail में समझते हैं।",
        "common_mistake_prefix": "एक important बात याद रखना।",
        "real_life_prefix": "एक real life example से समझो।",
        "numerical_concepts_intro": "पहले हम कुछ important concepts समझते हैं जो इस page पर use हुए हैं।",
        "to_find_prefix": "हमें ढूंढना है:",
        "final_answer_prefix": "Final answer:",
    },
    "hinglish": {
        "concepts_intro": "Pehle kuch important concepts samajhte hain jo is page pe hain.",
        "concept_detail_intro": "Ab hum ek ek concept ko detail mein samajhte hain.",
        "common_mistake_prefix": "Ek important baat yaad rakhna.",
        "real_life_prefix": "Ek real life example se samjho.",
        "numerical_concepts_intro": "Pehle hum kuch important concepts samjhte hain jo is page pe use hue hain.",
        "to_find_prefix": "Humhe dhundhna hai:",
        "final_answer_prefix": "Final answer:",
    },
    "gujarati": {
        "concepts_intro": "પહેલા આ page પર રહેલા કેટલાક મહત્વના concepts સમજીએ.",
        "concept_detail_intro": "હવે આપણે દરેક concept ને detail માં સમજીએ.",
        "common_mistake_prefix": "એક મહત્વની વાત યાદ રાખજો.",
        "real_life_prefix": "એક real life example થી સમજો.",
        "numerical_concepts_intro": "પહેલા આ page પર વપરાયેલા કેટલાક important concepts સમજીએ.",
        "to_find_prefix": "આપણે શોધવાનું છે:",
        "final_answer_prefix": "Final answer:",
    },
    "marathi": {
        "concepts_intro": "आधी या page वरचे काही महत्त्वाचे concepts समजून घेऊ.",
        "concept_detail_intro": "आता आपण प्रत्येक concept detail मध्ये समजून घेऊ.",
        "common_mistake_prefix": "एक महत्त्वाची गोष्ट लक्षात ठेवा.",
        "real_life_prefix": "एका real life example ने समजू.",
        "numerical_concepts_intro": "आधी या page वर वापरलेले काही important concepts समजून घेऊ.",
        "to_find_prefix": "आपल्याला शोधायचे आहे:",
        "final_answer_prefix": "Final answer:",
    },
    "bengali": {
        "concepts_intro": "প্রথমে এই page-এর কিছু গুরুত্বপূর্ণ concepts বুঝে নিই।",
        "concept_detail_intro": "এবার আমরা প্রতিটি concept বিস্তারিতভাবে বুঝি।",
        "common_mistake_prefix": "একটা গুরুত্বপূর্ণ কথা মনে রাখবে।",
        "real_life_prefix": "একটা real life example দিয়ে বুঝো।",
        "numerical_concepts_intro": "প্রথমে এই page-এ ব্যবহৃত কিছু important concepts বুঝে নিই।",
        "to_find_prefix": "আমাদের খুঁজতে হবে:",
        "final_answer_prefix": "Final answer:",
    },
    "tamil": {
        "concepts_intro": "முதலில் இந்த page-ல் உள்ள சில முக்கியமான concepts-ஐ புரிந்து கொள்வோம்.",
        "concept_detail_intro": "இப்போது ஒவ்வொரு concept-ஐயும் விரிவாக புரிந்து கொள்வோம்.",
        "common_mistake_prefix": "ஒரு முக்கியமான விஷயம் நினைவில் வைத்துக்கொள்ளுங்கள்.",
        "real_life_prefix": "ஒரு real life example மூலம் புரிந்து கொள்ளுங்கள்.",
        "numerical_concepts_intro": "முதலில் இந்த page-ல் பயன்படுத்தப்பட்ட சில important concepts-ஐ புரிந்து கொள்வோம்.",
        "to_find_prefix": "நாம் கண்டறிய வேண்டியது:",
        "final_answer_prefix": "Final answer:",
    },
    "telugu": {
        "concepts_intro": "ముందుగా ఈ page లో ఉన్న కొన్ని ముఖ్యమైన concepts అర్థం చేసుకుందాం.",
        "concept_detail_intro": "ఇప్పుడు ప్రతి concept ని వివరంగా అర్థం చేసుకుందాం.",
        "common_mistake_prefix": "ఒక ముఖ్యమైన విషయం గుర్తుంచుకోండి.",
        "real_life_prefix": "ఒక real life example తో అర్థం చేసుకోండి.",
        "numerical_concepts_intro": "ముందుగా ఈ page లో వాడిన కొన్ని important concepts అర్థం చేసుకుందాం.",
        "to_find_prefix": "మనం కనుగొనవలసింది:",
        "final_answer_prefix": "Final answer:",
    },
    "urdu": {
        "concepts_intro": "پہلے اس page پر موجود کچھ اہم concepts سمجھ لیتے ہیں۔",
        "concept_detail_intro": "اب ہم ہر concept کو تفصیل سے سمجھتے ہیں۔",
        "common_mistake_prefix": "ایک اہم بات یاد رکھیں۔",
        "real_life_prefix": "ایک real life example سے سمجھو۔",
        "numerical_concepts_intro": "پہلے ہم اس page پر استعمال ہونے والے کچھ important concepts سمجھتے ہیں۔",
        "to_find_prefix": "ہمیں ڈھونڈنا ہے:",
        "final_answer_prefix": "Final answer:",
    },
}

# Default fallback to English
_DEFAULT_PHRASES = TRANSITION_PHRASES["english"]

def _get_phrases(language):
    return TRANSITION_PHRASES.get(language.lower(), _DEFAULT_PHRASES)

LANGUAGE_VOICES = {
    "hindi":    "hi-IN-SwaraNeural",     # Male — change to hi-IN-SwaraNeural for female
    "hinglish": "hi-IN-SwaraNeural",     # Hinglish uses Hindi voice (handles code-switching)
    "gujarati": "gu-IN-NiranjanNeural",   # Male
    "marathi":  "mr-IN-AarohiNeural",     # Female
    "bengali":  "bn-IN-BashkarNeural",    # Male
    "tamil":    "ta-IN-ValluvarNeural",   # Male
    "telugu":   "te-IN-MohanNeural",      # Male    
    "urdu":     "ur-PK-AsadNeural",       # Male
    "english":  "en-US-GuyNeural",        # Male — change to en-US-JennyNeural for female
}

def _ensure_period(text):
    """Ensure text ends with sentence-ending punctuation so TTS adds a natural pause."""
    text = text.rstrip()
    if text and text[-1] not in ".!?।":
        text += "."
    return text


# TTS pause markers — edge_tts (Azure Neural TTS) interprets sentence-ending
# punctuation as natural pauses. Repeated periods create extra breathing room.
_PAUSE_BETWEEN_ITEMS    = " ... "    # ~0.7s gap between concepts / points
_PAUSE_BETWEEN_SECTIONS = " ...... " # ~1.2s gap between major sections (concepts → content → example)


def _generate_tts_file(text, voice, max_retries=3):
    """Generate TTS audio for a single text string and return the temp MP3 file path.
    Retries up to max_retries times on NoAudioReceived, then falls back to gTTS."""
    for attempt in range(max_retries):
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".mp3")
        loop = asyncio.new_event_loop()
        asyncio.set_event_loop(loop)
        try:
            loop.run_until_complete(
                edge_tts.Communicate(text=text, voice=voice).save(tmp.name)
            )
            return tmp.name
        except Exception:
            try:
                os.unlink(tmp.name)
            except OSError:
                pass
            if attempt < max_retries - 1:
                time.sleep(1.5)
        finally:
            loop.close()

    # Fallback: use gTTS when edge_tts keeps failing
    lang_code = LANGUAGE_CODES.get(
        next((k for k in LANGUAGE_CODES if LANGUAGE_VOICES.get(k) == voice), "english"),
        "en"
    )
    tmp_fallback = tempfile.NamedTemporaryFile(delete=False, suffix=".mp3")
    gTTS(text=text, lang=lang_code).save(tmp_fallback.name)
    return tmp_fallback.name


def _concat_audio_with_pause(paths, pause_ms=1500):
    """Concatenate MP3 audio files with a silent gap between each.
    Uses pydub for clean silence; falls back to raw byte concat if pydub is unavailable."""
    if len(paths) == 1:
        return paths[0]
    try:
        from pydub import AudioSegment
        silence = AudioSegment.silent(duration=pause_ms)
        combined = AudioSegment.empty()
        for i, path in enumerate(paths):
            seg = AudioSegment.from_mp3(path)
            combined += seg
            if i < len(paths) - 1:
                combined += silence
        out = tempfile.NamedTemporaryFile(delete=False, suffix=".mp3")
        combined.export(out.name, format="mp3")
        for p in paths:
            try:
                os.unlink(p)
            except OSError:
                pass
        return out.name
    except (ImportError, FileNotFoundError, OSError):
        # pydub not installed — concatenate raw MP3 bytes as fallback
        out = tempfile.NamedTemporaryFile(delete=False, suffix=".mp3")
        with open(out.name, "wb") as f:
            for path in paths:
                with open(path, "rb") as pf:
                    f.write(pf.read())
        for p in paths:
            try:
                os.unlink(p)
            except OSError:
                pass
        return out.name


def text_to_audio(state):
    task = state["task"]
    language = state["language"]
    voice = LANGUAGE_VOICES.get(language.lower(), "en-US-GuyNeural")
    audio_paths = []
    phrases = _get_phrases(language)

    def _tts(text):
        """Convert text to MP3, applying hinglish conversion when needed."""
        if language.lower() == "hinglish":
            text = hinglish_to_devanagari(text)
        return _generate_tts_file(text, voice)

    if task == "theory":
        for theory_response in state["explained_list"]:
            for slide in theory_response.slides:
                slide_segment_paths = []

                # 1. Concepts section — all concepts spoken as one audio clip
                if slide.concepts_used:
                    concept_parts = [_ensure_period(phrases["concepts_intro"])]
                    for concept in slide.concepts_used:
                        concept_parts.append(
                            _ensure_period(f"{concept.name}. {concept.concept_used_explanation}")
                        )
                    slide_segment_paths.append(_tts(_PAUSE_BETWEEN_ITEMS.join(concept_parts)))

                # 2. Content points — each point gets its OWN audio clip,
                #    then joined with 1.5s silence so the listener can absorb each point
                if slide.content_points:
                    point_paths = [
                        _tts(_ensure_period(pt.content_original_translation))
                        for pt in slide.content_points
                    ]
                    slide_segment_paths.append(
                        _concat_audio_with_pause(point_paths, pause_ms=1500)
                    )

                # 3. Real life example — one audio clip
                if slide.real_life_example:
                    slide_segment_paths.append(
                        _tts(_ensure_period(f"{phrases['real_life_prefix']} {slide.real_life_example}"))
                    )

                # Combine all three sections with 800ms pause between them
                if slide_segment_paths:
                    audio_paths.append(
                        _concat_audio_with_pause(slide_segment_paths, pause_ms=800)
                        if len(slide_segment_paths) > 1 else slide_segment_paths[0]
                    )

    else:
        # Maths — keep text-based assembly, one TTS call per slide
        texts_to_speak = []
        for item in state["maths_explained"]:
            sections = []
            item_type = item.get("type", "numerical") if isinstance(item, dict) else "numerical"

            if item_type == "conceptual":
                concept_response = item["data"]
                # Page summary
                sections.append(_ensure_period(concept_response.page_summary))
                # Each concept explained
                concept_parts = [_ensure_period(phrases["concept_detail_intro"])]
                for concept in concept_response.concepts:
                    concept_parts.append(
                        _ensure_period(f"{concept.name}. {concept.simple_explanation}")
                    )
                    if concept.when_to_use:
                        concept_parts.append(_ensure_period(concept.when_to_use))
                    if concept.common_mistakes:
                        concept_parts.append(
                            _ensure_period(f"{phrases['common_mistake_prefix']} {concept.common_mistakes}")
                        )
                sections.append(_PAUSE_BETWEEN_ITEMS.join(concept_parts))
            else:
                response_model = item["data"] if isinstance(item, dict) else item
                # 1. Concepts used section
                if response_model.concepts_used:
                    concept_parts = [_ensure_period(phrases["numerical_concepts_intro"])]
                    for concept in response_model.concepts_used:
                        concept_parts.append(
                            _ensure_period(f"{concept.name}. {concept.explanation}")
                        )
                    sections.append(_PAUSE_BETWEEN_ITEMS.join(concept_parts))

                # 2. Each question
                for question in response_model.questions:
                    q_parts = []
                    q_parts.append(
                        _ensure_period(f"Question {question.question_number}. {question.original_question_text}")
                    )
                    q_parts.append(_ensure_period(question.original_question_text_explanation))
                    q_parts.append(
                        _ensure_period(f"{phrases['to_find_prefix']} {question.to_find}")
                    )
                    q_parts.append(_ensure_period(question.to_find_explanation))
                    for step in question.steps:
                        q_parts.append(_ensure_period(step.simple_explanation))
                    q_parts.append(
                        _ensure_period(f"{phrases['final_answer_prefix']} {question.final_answer}")
                    )
                    sections.append(_PAUSE_BETWEEN_ITEMS.join(q_parts))

            texts_to_speak.append(_PAUSE_BETWEEN_SECTIONS.join(sections))

        for text in texts_to_speak:
            audio_paths.append(_tts(text))

    return {"audio_list": audio_paths}   


def answer_question_direct(explanation_text, user_query, language):
    has_context = bool((explanation_text or "").strip())

    if has_context:
        qa_prompt = PromptTemplate(
            template="""You are a friendly classmate who just finished reading this slide and your friend is asking you a question about it — answer like you're texting them, not writing an exam answer.

Here is the slide content:
{slide_content}

The student's question is:
{question}

Answer in {language}.

IMPORTANT — Language style rules:
- Write exactly how real people talk in everyday life, NOT like a textbook or newspaper.
- Use short sentences. Keep it simple and warm.
- If you are writing in Hindi, write like a friend would WhatsApp you — not like a Hindi textbook. For example: say "मतलब ये है कि..." not "इसका तात्पर्य यह है कि..."; say "काम करता है" not "कार्य संपादित करता है".
- It is totally fine — and often better — to keep common English words (like "basically", "simple", "process", "point") if that is how people naturally say it in that language. Do not force a heavy translation when a mixed word feels more natural.
- Never use a word that would make a normal person stop and think "what does that mean?"
- If the question is not related to the slide content at all, just say politely in {language} that "yeh slide mein nahi hai, lekin..." and try to help briefly from general knowledge.

How to answer:
- Get straight to the point — no long intros.
- If it needs an example to click, give one short real-life example.
- Max 4-5 lines. Warm, clear, and confidence-building.
- End with something encouraging like "clear hua?" or its equivalent in {language}.""",
            input_variables=["slide_content", "question", "language"]
        )
        chain = qa_prompt | model_text | parser
        answer = chain.invoke({
            "slide_content": explanation_text,
            "question": user_query,
            "language": language,
        })
    else:
        qa_prompt = PromptTemplate(
            template="""You are a helpful AI tutor and problem solver.

The student's question is:
{question}

Answer in {language}.

Rules:
- Give a direct, practical answer.
- Keep language natural and easy to understand.
- Use short sentences and avoid heavy jargon.
- If needed, give one short example.
- If the question is ambiguous, state one reasonable assumption and answer.
- Max 6 lines.
- End with a helpful follow-up like: "Want a quick step-by-step too?" in {language}.""",
            input_variables=["question", "language"]
        )
        chain = qa_prompt | model_text | parser
        answer = chain.invoke({
            "question": user_query,
            "language": language,
        })

    voice = LANGUAGE_VOICES.get(language.lower(), "en-US-GuyNeural")
    # For hinglish: convert romanized Hindi to Devanagari before TTS
    tts_text = hinglish_to_devanagari(answer) if language.lower() == "hinglish" else answer
    tmp   = tempfile.NamedTemporaryFile(delete=False, suffix=".mp3")
    loop  = asyncio.new_event_loop()
    asyncio.set_event_loop(loop)
    try:
        loop.run_until_complete(
            edge_tts.Communicate(text=tts_text, voice=voice).save(tmp.name)
        )
    finally:
        loop.close()

    return {
        "query_ans":   answer,
        "query_audio": tmp.name
    }
    
from langchain_core.output_parsers import PydanticOutputParser
from langchain_core.prompts import PromptTemplate
from pydantic import BaseModel, Field
from typing import List

class particular_component(BaseModel):
    question: str = Field(description="question on the basis of context")
    ans: str = Field(description="solution for the given question")

class multi_component(BaseModel):
    multi_questions: List[particular_component]

question_parser = PydanticOutputParser(pydantic_object=multi_component)

def question_generator(all_documents, task="theory"):
    
    if task == "maths":
        question_prompt = PromptTemplate(
            template="""
You are an expert mathematics question setter and tutor.

Your task is to generate EXACTLY 2 NEW math practice questions along with their detailed step-by-step solutions.

⚠️ CRITICAL RULE — CHANGE THE VALUES:
- Study the mathematical concepts, formulas, and problem-solving methods shown in the context.
- Then create BRAND NEW questions that use the SAME concepts and logic but with COMPLETELY DIFFERENT numerical values, variable names, or problem setups.
- DO NOT copy or reuse any question from the context as-is.
- The new questions must be solvable using the same methods/formulas demonstrated in the context.
- Keep the difficulty level similar to the context.

Instructions:
- Generate exactly 2 questions (not more, not less).
- Each question must use different numbers/values than those in the context.
- Each answer must show the full step-by-step solution with the new values.
- The underlying concept/formula must match what is taught in the context.

Context (study the methods and concepts here, but DO NOT reuse the same numbers):
{context}

Output Format (STRICTLY follow this JSON structure):
{format_instructions}
""",
            input_variables=["context"],
            partial_variables={
                "format_instructions": question_parser.get_format_instructions()
            }
        )
    else:
        question_prompt = PromptTemplate(
            template="""
You are an expert question setter and tutor.

Your task is to generate EXACTLY 2 high-quality questions along with their detailed solutions based ONLY on the given context.

Instructions:
- Generate exactly 2 questions (not more, not less).
- Questions should test understanding, not just direct copying.
- Avoid very easy or trivial questions.
- Each question must be clear, concise, and meaningful.
- Each answer must be well-explained, step-by-step if required.
- Do NOT use any external knowledge. Stick strictly to the context.

Context:
{context}

Output Format (STRICTLY follow this JSON structure):
{format_instructions}
""",
            input_variables=["context"],
            partial_variables={
                "format_instructions": question_parser.get_format_instructions()
            }
        )

    chain = question_prompt | model_text | question_parser

    all_results = []

    for doc in all_documents:
        
        if not doc["doc"] or len(doc["doc"].strip()) < 50:
            continue
        
        try:
            result = chain.invoke({"context": doc["doc"]})
            all_results.append(result)
        except Exception:
            continue

    return all_results

class option(BaseModel):
    single_option:str=Field(description="this is the option for the question")
    
class single_que(BaseModel):
    que:str=Field(description="the question")
    ans:str=Field(description="the answer")
    options:List[option]=Field(description="all four options")
    
class quiz(BaseModel):
    full_ques:List[single_que]

quiz_parser=PydanticOutputParser(pydantic_object=quiz)

def quiz_generator(all_docs, task="theory"):
    ques_list=[]
    for doc in all_docs:
        context=doc["doc"]
        if task == "maths":
            prompt = PromptTemplate(
    template="""
You are an expert mathematics exam paper setter.

Your task is to generate HIGH-QUALITY multiple choice questions (MCQs) that test the SAME mathematical concepts shown in the context but with COMPLETELY DIFFERENT numerical values.

========================
🎯 INSTRUCTIONS
========================

- Generate EXACTLY 2 MCQs.
- ⚠️ CRITICAL: Study the math concepts/formulas in the context, then create NEW questions using DIFFERENT numbers, values, or variable setups.
- DO NOT copy any question from the context as-is. Change ALL numerical values.
- The questions must be solvable using the same methods/formulas from the context.
- Keep difficulty level similar to the context.
- Each question must have EXACTLY 4 options.
- Only ONE option should be correct (computed with the new values).
- The correct answer must be included in the options.
- All wrong options should be plausible (common mistakes with the new values).

========================
📘 CONTEXT (learn the concepts here, but use DIFFERENT numbers)
========================
{context}

========================
📦 OUTPUT FORMAT (STRICT)
========================
{format_instructions}

========================
⚠️ IMPORTANT RULES
========================
- DO NOT reuse the same numbers from the context.
- Ensure the correct answer is freshly computed for the new values.
- Ensure options are realistic and relevant.
- Ensure no duplicate options.
- Ensure answer matches one of the options EXACTLY.
""",
    input_variables=["context"],
    partial_variables={
        "format_instructions": quiz_parser.get_format_instructions()
    }
)
        else:
            prompt = PromptTemplate(
    template="""
You are an expert exam paper setter.

Your task is to generate HIGH-QUALITY multiple choice questions (MCQs) based ONLY on the given context.

========================
🎯 INSTRUCTIONS
========================

- Generate EXACTLY 2 MCQs from the context.
- Each question must test understanding, not simple copying.
- Each question must have EXACTLY 4 options.
- Only ONE option should be correct.
- The correct answer must be included in the options.
- Do NOT mention which option is correct inside the options.
- Avoid vague or ambiguous questions.
- Questions should be clear, concise, and meaningful.

========================
📘 CONTEXT
========================
{context}

========================
📦 OUTPUT FORMAT (STRICT)
========================
{format_instructions}

========================
⚠️ IMPORTANT RULES
========================
- Do NOT use any external knowledge.
- Do NOT generate more than 2 questions.
- Ensure options are realistic and relevant.
- Ensure no duplicate options.
- Ensure answer matches one of the options EXACTLY.
""",
    input_variables=["context"],
    partial_variables={
        "format_instructions": quiz_parser.get_format_instructions()
    }
)
        chain=prompt|model_text|quiz_parser
        result=_invoke_with_retry(chain, {"context":context})
        ques_list.append(result)
        
    return  ques_list
        
from docx import Document as DocxDoc
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io

from docx import Document as DocxDoc
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io

def generate_quiz_docx(quiz_data, doc_filename="Document"):
    """
    Generates a .docx quiz document from quiz_data.
    Questions first, Answer Key at the bottom.
    Returns bytes (file content).
    """
    doc = DocxDoc()

    # ── Page margins ──
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)

    # ── Title ──
    title = doc.add_heading("Quiz — Practice Questions", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.color.rgb = RGBColor(0x1d, 0x7f, 0xe8)

    sub = doc.add_paragraph(f"Document: {doc_filename}")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size  = Pt(10)
    sub.runs[0].font.color.rgb = RGBColor(0x6b, 0x87, 0xa8)

    doc.add_paragraph("")  # spacer

    option_labels = ["A", "B", "C", "D"]

    # ── SECTION 1: Questions ──
    q_heading = doc.add_heading("Questions", level=1)
    q_heading.runs[0].font.color.rgb = RGBColor(0x1d, 0x7f, 0xe8)

    for i, q in enumerate(quiz_data, start=1):

        # Question text
        q_para = doc.add_paragraph()
        q_para.paragraph_format.space_before = Pt(8)
        q_run = q_para.add_run(f"Q{i}. {q['question']}")
        q_run.bold = True
        q_run.font.size = Pt(11)

        # Options
        for j, opt in enumerate(q["options"]):
            label = option_labels[j] if j < len(option_labels) else str(j + 1)
            opt_para = doc.add_paragraph()
            opt_para.paragraph_format.left_indent = Inches(0.4)
            opt_para.paragraph_format.space_after = Pt(2)
            opt_run = opt_para.add_run(f"{label}. {opt}")
            opt_run.font.size = Pt(10.5)

        doc.add_paragraph("")  # spacer

    # ── Page break before answer key ──
    doc.add_page_break()

    # ── SECTION 2: Answer Key ──
    ans_heading = doc.add_heading("Answer Key", level=1)
    ans_heading.runs[0].font.color.rgb = RGBColor(0x22, 0xc5, 0x5e)

    ans_intro = doc.add_paragraph(
        "Answers are listed below. Review after attempting all questions."
    )
    ans_intro.runs[0].font.size  = Pt(10)
    ans_intro.runs[0].font.color.rgb = RGBColor(0x6b, 0x87, 0xa8)
    ans_intro.paragraph_format.space_after = Pt(12)

    for i, q in enumerate(quiz_data, start=1):
        correct = q["answer"]
        options = q["options"]

        # Find correct option label
        correct_label = "—"
        for j, opt in enumerate(options):
            if opt.strip().lower() == correct.strip().lower():
                correct_label = option_labels[j] if j < len(option_labels) else str(j + 1)
                break

        ans_para = doc.add_paragraph()
        ans_para.paragraph_format.space_after = Pt(4)

        num_run = ans_para.add_run(f"Q{i}. ")
        num_run.bold = True
        num_run.font.size = Pt(11)

        lbl_run = ans_para.add_run(f"({correct_label}) ")
        lbl_run.bold = True
        lbl_run.font.size = Pt(11)
        lbl_run.font.color.rgb = RGBColor(0x22, 0xc5, 0x5e)

        ans_run = ans_para.add_run(correct)
        ans_run.font.size = Pt(10.5)
        ans_run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)

    # ── Save to bytes ──
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

import io
from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def generate_qa_docx(qa_bank: list, doc_title: str = "Practice Questions") -> bytes:
    """
    Generate a .docx from a list of {page_no, question, answer} dicts.
    Returns raw bytes of the .docx file.
    Format: Q1) ...  Ans: ...
    """
    doc = DocxDocument()

    # ── Title ──
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(doc_title)
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x1D, 0x7F, 0xE8)  # brand blue

    doc.add_paragraph()  # spacer

    # Group by page
    from collections import defaultdict
    page_groups = defaultdict(list)
    for qa in qa_bank:
        page_groups[qa["page_no"]].append(qa)

    global_q_no = 1
    for pg_no in sorted(page_groups.keys()):
        display_pg = pg_no + 1

        # ── Page heading ──
        pg_para = doc.add_paragraph()
        pg_run = pg_para.add_run(f"Page {display_pg}")
        pg_run.bold = True
        pg_run.font.size = Pt(13)
        pg_run.font.color.rgb = RGBColor(0x3B, 0x9E, 0xFF)

        for qa in page_groups[pg_no]:
            # ── Question ──
            q_para = doc.add_paragraph()
            q_run = q_para.add_run(f"Q{global_q_no}) {qa['question']}")
            q_run.bold = True
            q_run.font.size = Pt(11)

            # ── Answer ──
            a_para = doc.add_paragraph()
            a_para.paragraph_format.left_indent = Pt(20)
            a_run = a_para.add_run(f"Ans: {qa['answer']}")
            a_run.font.size = Pt(11)
            a_run.font.color.rgb = RGBColor(0x22, 0x8B, 0x22)

            doc.add_paragraph()  # spacer between Q&A pairs
            global_q_no += 1

    # Serialize to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


graph = StateGraph(state)

graph.add_node("extract_particular", extract_particular)
graph.add_node("single_multi_page", single_multi_page)
graph.add_node("text_to_audio", text_to_audio)  


graph.add_edge(START, "extract_particular")
graph.add_edge("extract_particular", "single_multi_page")
graph.add_edge("single_multi_page", "text_to_audio")  


graph.add_edge("text_to_audio", END)
workflow = graph.compile()